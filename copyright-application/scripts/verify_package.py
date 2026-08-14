#!/usr/bin/env python3
from __future__ import annotations

import json
import re
import sys
from pathlib import Path

from docx import Document
from pypdf import PdfReader

ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "copyright-application"
QA = OUT / "qa"
DATA = json.loads((OUT / "application-data.yaml").read_text(encoding="utf-8"))
SELECTION = json.loads((QA / "source-selection.json").read_text(encoding="utf-8"))
REQUIRED = [
    OUT / "01-软件著作权登记信息表.docx",
    OUT / "02-Blog and OJ V1.0软件操作说明书.docx",
    OUT / "02-Blog and OJ V1.0软件操作说明书.pdf",
    OUT / "03-Blog and OJ V1.0源程序鉴别材料.docx",
    OUT / "03-Blog and OJ V1.0源程序鉴别材料.pdf",
    OUT / "04-申请材料清单与提交指南.docx",
    OUT / "05-待填写信息清单.md",
]
SENSITIVE = {
    "PRC identity number": re.compile(r"(?<!\d)\d{17}[0-9Xx](?!\d)"),
    "mainland mobile number": re.compile(r"(?<!\d)1[3-9]\d{9}(?!\d)"),
    "private key": re.compile(r"-----BEGIN (?:RSA |EC |OPENSSH )?PRIVATE KEY-----"),
    "JWT-like token": re.compile(r"\beyJ[A-Za-z0-9_-]{20,}\.[A-Za-z0-9_-]{10,}\.[A-Za-z0-9_-]{10,}\b"),
    "cloud access key": re.compile(r"\b(?:AKIA|LTAI)[A-Z0-9]{12,}\b"),
}


def docx_text(path: Path) -> str:
    doc = Document(path)
    chunks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    for section in doc.sections:
        chunks.extend(p.text for p in section.header.paragraphs)
        chunks.extend(p.text for p in section.footer.paragraphs)
    return "\n".join(chunks)


def pdf_text(path: Path) -> tuple[int, str]:
    reader = PdfReader(path)
    return len(reader.pages), "\n".join(page.extract_text() or "" for page in reader.pages)


def main() -> None:
    checks: list[tuple[str, bool, str]] = []
    for path in REQUIRED:
        checks.append((f"存在 {path.name}", path.exists() and path.stat().st_size > 0, str(path.stat().st_size if path.exists() else 0)))

    info = docx_text(OUT / "01-软件著作权登记信息表.docx")
    manual_pages, manual = pdf_text(OUT / "02-Blog and OJ V1.0软件操作说明书.pdf")
    source_pages, source = pdf_text(OUT / "03-Blog and OJ V1.0源程序鉴别材料.pdf")
    guide = docx_text(OUT / "04-申请材料清单与提交指南.docx")
    pending = (OUT / "05-待填写信息清单.md").read_text(encoding="utf-8")
    combined = "\n".join((info, manual, source, guide, pending))

    checks.extend([
        ("说明书页数不少于15", manual_pages >= 15, str(manual_pages)),
        ("源程序PDF恰为60页", source_pages == 60, str(source_pages)),
        ("选择清单恰为3000行", len(SELECTION["lines"]) == 3000, str(len(SELECTION["lines"]))),
        ("每页恰为50行", all(sum(1 for r in SELECTION["lines"] if r["page"] == p) == 50 for p in range(1, 61)), "pages 1-60"),
        ("页内行号连续", all(r["page_line"] == (i - 1) % 50 + 1 for i, r in enumerate(SELECTION["lines"], 1)), "1-50 repeated"),
        ("总行号连续", all(r["deposit_line"] == i for i, r in enumerate(SELECTION["lines"], 1)), "1-3000"),
    ])

    facts = ["Blog and OJ", "V1.0", "2026年8月10日", "2026年2月23日"]
    compact_info = re.sub(r"\s+", "", info)
    compact_manual = re.sub(r"\s+", "", manual)
    for fact in facts:
        compact_fact = re.sub(r"\s+", "", fact)
        checks.append((f"信息表包含 {fact}", compact_fact in compact_info, fact))
        checks.append((f"说明书包含 {fact}", compact_fact in compact_manual, fact))
    checks.append(("提交指南包含软件名称", "Blog and OJ" in guide and "V1.0" in guide, "Blog and OJ V1.0"))

    for label, pattern in SENSITIVE.items():
        checks.append((f"未发现 {label}", not pattern.search(combined), "no match"))

    placeholders = ("TBD", "TODO", "待补截图", "PLACEHOLDER")
    checks.append(("无内部占位标记", not any(word in combined for word in placeholders), ", ".join(placeholders)))

    failures = [item for item in checks if not item[1]]
    report = ["# Blog and OJ V1.0 申请材料验证报告", "", f"结果：{'PASS' if not failures else 'FAIL'}", ""]
    report.extend(f"- {'通过' if ok else '失败'}：{name}（{detail}）" for name, ok, detail in checks)
    report.append("")
    report.append("说明：个人姓名、证件号码、地址、电话和邮箱按用户要求保持空白。")
    (QA / "verification-report.md").write_text("\n".join(report) + "\n", encoding="utf-8")
    if failures:
        for name, _, detail in failures:
            print(f"FAIL: {name}: {detail}")
        raise SystemExit(1)
    print(f"PASS: {len(checks)} checks")
    print((QA / "verification-report.md").relative_to(ROOT))


if __name__ == "__main__":
    main()
