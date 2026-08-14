#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "copyright-application"
DATA = json.loads((OUT / "application-data.yaml").read_text(encoding="utf-8"))
BLUE = "2E74B5"
LIGHT_BLUE = "E8EEF5"
GRAY = "F2F4F7"
MUTED = RGBColor(90, 100, 112)
BLANK = "________________"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def style_table(table, widths=(2700, 6660)) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row_idx, row in enumerate(table.rows):
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.15
                for run in paragraph.runs:
                    set_font(run, 10.5)
            if idx == 0:
                set_cell_shading(cell, LIGHT_BLUE)
                for run in cell.paragraphs[0].runs:
                    run.bold = True


def set_font(run, size=11, bold=None, color=None, name="STHeiti") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def configure_document(doc: Document, title: str) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.4)
    normal = doc.styles["Normal"]
    normal.font.name = "STHeiti"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "STHeiti")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "STHeiti")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "STHeiti")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for style_name, size, before, after, color in (
        ("Title", 24, 0, 12, BLUE),
        ("Heading 1", 16, 18, 10, BLUE),
        ("Heading 2", 13, 14, 7, BLUE),
        ("Heading 3", 12, 10, 5, "1F4D78"),
    ):
        style = doc.styles[style_name]
        style.font.name = "STHeiti"
        style._element.rPr.rFonts.set(qn("w:ascii"), "STHeiti")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "STHeiti")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "STHeiti")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = style_name != "Title"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    header = section.header.paragraphs[0]
    header.text = title
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header.runs[0], 8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Blog and OJ V1.0")
    set_font(run, 8.5, color=MUTED)
    doc.core_properties.title = title
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.space_after = Pt(8)
    set_font(p.add_run(title), 24, bold=True, color=RGBColor.from_string(BLUE))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    set_font(p.add_run(subtitle), 13, color=MUTED)


def add_field_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value if value else BLANK
    style_table(table)


def bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    set_font(p.add_run(text), 11)


def build_info_form() -> Path:
    software, applicant, tech = DATA["software"], DATA["applicant"], DATA["technical"]
    doc = Document()
    configure_document(doc, "软件著作权登记信息表")
    add_title(doc, "软件著作权登记信息表", "Blog and OJ V1.0 · 个人申请草稿")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("说明：本表用于整理在线申请字段，带横线内容由申请人提交前补充。"), 10, color=MUTED)
    doc.add_heading("一、软件基本信息", level=1)
    add_field_table(doc, [
        ("软件全称", software["full_name"]), ("软件简称", software["short_name"]),
        ("版本号", software["version"]), ("软件分类", software["category"]),
        ("开发方式", software["development_type"]), ("权利取得方式", software["rights_acquisition"]),
        ("权利范围", software["rights_scope"]), ("开发完成日期", software["completion_date"]),
    ])
    doc.add_heading("二、发表信息", level=1)
    add_field_table(doc, [
        ("发表状态", "已发表"), ("首次发表日期", software["publication_date"]),
        ("首次发表地点", software["publication_place"]), ("首次发表方式", software["publication_method"]),
        ("发布地址", software["publication_url"]),
    ])
    doc.add_heading("三、著作权人信息", level=1)
    add_field_table(doc, [
        ("申请人类型", applicant["type"]), ("姓名", applicant["name"]),
        ("证件类型", applicant["id_type"]), ("证件号码", applicant["id_number"]),
        ("联系地址", applicant["address"]), ("邮政编码", applicant["postal_code"]),
        ("联系电话", applicant["phone"]), ("电子邮箱", applicant["email"]),
    ])
    doc.add_heading("四、软件技术信息", level=1)
    add_field_table(doc, [
        ("开发目的", tech["purpose"]), ("所属领域/行业", tech["industry"]),
        ("硬件环境", tech["hardware_environment"]), ("软件环境", tech["software_environment"]),
        ("开发工具", tech["development_tools"]), ("编程语言", tech["languages"]),
        ("源程序量", tech["source_scale"]),
    ])
    doc.add_heading("五、主要功能和技术特点", level=1)
    for item in tech["features"]:
        bullet(doc, item)
    doc.add_heading("六、申请人确认", level=1)
    doc.add_paragraph("本人确认以上信息真实、准确，并确认该软件为本人独立开发且依法享有软件著作权。")
    add_field_table(doc, [("申请人签名", ""), ("确认日期", "____年____月____日")])
    path = OUT / "01-软件著作权登记信息表.docx"
    doc.save(path)
    return path


def build_guide() -> Path:
    doc = Document()
    configure_document(doc, "申请材料清单与提交指南")
    add_title(doc, "申请材料清单与提交指南", "Blog and OJ V1.0 · 个人申请")
    doc.add_heading("一、提交材料清单", level=1)
    for text in (
        "软件著作权登记在线申请信息；",
        "Blog and OJ V1.0源程序鉴别材料（普通交存，共60页）；",
        "Blog and OJ V1.0软件操作说明书；",
        "申请人有效身份证明；",
        "登记机构要求补充的其他权属证明（仅在实际存在委托、合作、职务或继受情形时提供）。",
    ):
        bullet(doc, text)
    doc.add_heading("二、建议提交顺序", level=1)
    for idx, text in enumerate((
        "登录中国版权保护中心软件登记系统，完成个人实名认证。",
        "参照《软件著作权登记信息表》逐项填写在线申请信息。",
        "上传源程序鉴别材料和软件操作说明书，并核对文件名称、软件名称和版本号。",
        "按系统要求上传身份证明并完成确认或签章。",
        "提交后保存受理凭证，留意补正通知并在规定期限内处理。",
    ), 1):
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.5)
        set_font(p.add_run(text), 11)
    doc.add_heading("三、提交前核查", level=1)
    checks = (
        "所有文件中的软件名称均为“Blog and OJ”，版本号均为“V1.0”。",
        "开发完成日期为2026年8月10日，首次发表日期为2026年2月23日。",
        "著作权人信息与身份证明完全一致，地址和联系方式可正常联系。",
        "源程序材料共60页，每页不少于50行，页眉和页码连续。",
        "说明书中的功能、界面和操作步骤与实际软件一致。",
        "材料中没有密码、密钥、令牌、真实数据库连接或其他不宜公开的信息。",
        "GitHub仓库首次公开日期和个人独立开发的权属事实已经再次核实。",
    )
    for check in checks:
        bullet(doc, "□ " + check)
    doc.add_heading("四、重要提示", level=1)
    doc.add_paragraph("本材料包用于辅助整理登记资料，不代替中国版权保护中心在线申请表，也不构成登记必然获批的承诺。在线系统字段、上传方式或审查要求如有更新，应以提交当日系统提示为准。")
    path = OUT / "04-申请材料清单与提交指南.docx"
    doc.save(path)
    return path


def build_pending() -> Path:
    text = """# Blog and OJ V1.0 待填写信息清单

提交申请前，请在线下或中国版权保护中心申请系统内补充以下信息。为保护隐私，不要把填写后的证件号码或联系方式提交到本代码仓库。

- 申请人姓名：须与身份证明完全一致。
- 证件号码：填写本人有效居民身份证号码。
- 联系地址：填写能够接收通知或材料的常用地址。
- 邮政编码：与联系地址对应。
- 联系电话：填写可正常接听的本人手机号。
- 电子邮箱：填写可正常接收通知的常用邮箱。
- 申请人签名或在线确认：按提交系统当前要求办理。
- 最终申请日期：以实际提交日期为准。

提交前还应再次确认：GitHub仓库于2026年2月23日已经公开；Blog and OJ V1.0由申请人独立开发，不属于职务、委托或合作开发。
"""
    path = OUT / "05-待填写信息清单.md"
    path.write_text(text, encoding="utf-8")
    return path


def add_manual_page(doc: Document, number: int, title: str, intro: str,
                    sections: list[tuple[str, list[str]]], image: Path | None = None,
                    caption: str | None = None) -> None:
    if number > 1:
        doc.add_page_break()
    doc.add_heading(f"{number}. {title}", level=1)
    p = doc.add_paragraph(intro)
    p.paragraph_format.space_after = Pt(10)
    if image and image.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(image), width=Inches(6.25))
        cp = doc.add_paragraph(caption or title)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(cp.add_run(""), 9, color=MUTED)
        cp.paragraph_format.space_after = Pt(10)
    for heading, items in sections:
        doc.add_heading(heading, level=2)
        for item in items:
            bullet(doc, item)


def build_manual() -> Path:
    shots = OUT / "assets" / "screenshots"
    doc = Document()
    configure_document(doc, "Blog and OJ V1.0 软件操作说明书")
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    set_font(p.add_run("BLOG AND OJ"), 12, bold=True, color=MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    set_font(p.add_run("软件操作说明书"), 30, bold=True, color=RGBColor.from_string(BLUE))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("版本 V1.0"), 16, color=MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    set_font(p.add_run("开发完成日期：2026年8月10日"), 11)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("首次发表日期：2026年2月23日"), 11)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("个人独立开发 · 软件著作权登记材料"), 10, color=MUTED)

    pages = [
        ("文档说明", "本说明书依据Blog and OJ V1.0当前源代码、项目说明和本地可运行界面编写，用于说明软件组成、运行条件和主要操作。",
         [("阅读对象", ["软件著作权登记审查人员及软件使用者。", "文中所有功能均可在当前仓库找到对应页面、接口或业务实现。"]),
          ("术语", ["Blog指博客内容模块；OJ指Online Judge在线判题模块。", "管理员指具有administrator身份标识并通过路由权限检查的用户。"])], None, None),
        ("软件概述", "Blog and OJ将个人博客与在线判题功能集成在一个前后端分离的网站中。",
         [("建设目的", ["为个人内容创作、知识分享和在线编程练习提供统一入口。", "将文章管理、题库、代码提交、判题反馈和学习记录形成完整闭环。"]),
          ("主要用户", ["普通用户可以阅读和管理博客、浏览题目、提交Java代码并查询记录。", "管理员可以维护用户、登录记录、题目和测试用例。"])], None, None),
        ("系统组成", "系统由Vue前端、Spring Boot主业务服务和独立判题服务组成。",
         [("前端Web", ["提供导航、认证、博客、题库、代码编辑、提交记录和管理页面。"]),
          ("主业务服务", ["处理身份认证、博客和题目数据、提交记录、对象存储及WebSocket连接。"]),
          ("判题服务", ["通过RabbitMQ接收任务，在Docker容器中编译运行Java代码并返回判题结果。"])], None, None),
        ("运行环境", "部署和使用软件前应准备相应服务端和客户端环境。",
         [("服务端", ["JDK 8和Maven 3.8及以上版本。", "MySQL 8、Redis、RabbitMQ、Docker和MinIO。", "建议4GB及以上内存、10GB及以上可用磁盘空间。"]),
          ("客户端", ["使用支持ECMAScript 2015及以上标准的现代桌面浏览器。", "浏览器需要允许JavaScript和WebSocket连接。 "])], None, None),
        ("安装与启动", "系统采用前后端分离方式启动，基础设施和业务服务需要按依赖顺序运行。",
         [("启动顺序", ["先启动MySQL、Redis、RabbitMQ、Docker和MinIO。", "再启动backend主业务服务和evaluatesystem判题服务。", "最后在web目录执行npm run serve并访问前端地址。"]),
          ("配置原则", ["数据库、消息队列、对象存储和Redis连接信息通过环境变量提供。", "真实账号、密码和访问密钥不得写入代码仓库或申请材料。 "])], None, None),
        ("用户登录", "未登录用户访问受保护页面时，系统跳转到登录页面并保留原访问路径。",
         [("操作步骤", ["在导航栏选择“登录”。", "输入用户名和密码。", "单击“登录”，系统校验账号状态并获取JWT身份凭证。"]),
          ("结果", ["登录成功后进入原目标页面或首页。", "业务错误、认证错误或连接错误会显示对应提示，不展示服务器敏感信息。 "])], shots / "01-login.png", "图1 用户登录页面"),
        ("用户注册", "新用户可以通过注册页面建立平台账号。",
         [("操作步骤", ["在导航栏选择“注册”。", "输入用户名、密码和确认密码。", "确认两次密码一致后单击“注册”。"]),
          ("校验", ["系统检查必填项、用户名可用性和两次密码一致性。", "注册成功后可返回登录页面使用新账号登录。 "])], shots / "02-register.png", "图2 用户注册页面"),
        ("博客浏览", "登录用户可以在博客页面分页浏览和搜索已发布内容。",
         [("操作步骤", ["从导航栏进入“博客”。", "输入关键词后执行搜索，系统从后端检索完整博客集合并返回当前页。", "单击博客条目进入详情页面，支持键盘焦点和回车操作。"]),
          ("显示内容", ["列表展示标题、作者、修改时间等摘要信息。", "详情页解析Markdown内容并对不安全HTML和链接进行过滤。 "])], None, None),
        ("博客创建与编辑", "用户可以维护本人博客，并在正文中插入和调整图片。",
         [("新建", ["进入个人博客管理页面并选择新建。", "填写标题、描述和正文，首次保存后获得博客编号。", "需要图片时选择本地图片，系统调整尺寸后上传到MinIO。"]),
          ("编辑与删除", ["选择本人博客进入编辑页面，修改内容后保存。", "删除操作通过与博客编号绑定的确认弹窗执行，避免误删其他条目。 "])], None, None),
        ("题库浏览", "题库页面按题号分页显示可练习题目，并支持全库关键词搜索。",
         [("操作步骤", ["从导航栏进入“题库”。", "通过分页控件切换当前页，或输入关键词后从第一页开始搜索。", "单击题目整行进入题目详情。"]),
          ("题目内容", ["详情包含标题、题目描述、输入输出说明、时间限制、内存限制和示例。 "])], None, None),
        ("代码编辑与提交", "题目详情页提供Ace Editor代码编辑区域，当前版本支持提交Java代码。",
         [("操作步骤", ["阅读题目要求和输入输出格式。", "在代码编辑器中完成Java程序。", "单击提交后等待服务端接受任务，提交按钮显示等待状态。"]),
          ("异常处理", ["HTTP提交失败时显示错误信息并结束等待。", "HTTP成功仅表示任务已接收，界面继续等待WebSocket判题结果。 "])], None, None),
        ("异步判题流程", "主业务服务与判题服务通过RabbitMQ解耦，用户代码在Docker容器中隔离执行。",
         [("处理流程", ["主业务服务校验用户和题目后保存提交记录。", "判题任务写入RabbitMQ队列，判题服务消费任务。", "判题服务创建Docker容器，编译并运行Java代码，收集所有标准输出和错误输出帧。", "判题结果经消息队列返回并通过WebSocket推送到浏览器。"]),
          ("结果类型", ["结果包括接受、答案错误、编译错误、运行错误、超时和系统错误等状态。 "])], None, None),
        ("提交记录与答案详情", "用户可以查询个人历史提交并查看单次提交的状态、代码和题目关联信息。",
         [("操作步骤", ["从导航栏进入提交记录。", "在列表中查看提交时间、题目、语言、得分和状态。", "选择记录进入详情，核对提交代码和判题反馈。"]),
          ("数据原则", ["列表只展示当前登录用户有权查看的记录。", "错误状态不会伪造得分或答案详情入口。 "])], None, None),
        ("后台管理", "管理员身份通过后端用户信息和前端路由守卫共同判断。",
         [("用户与登录记录", ["管理员可查询用户状态并执行封禁或解封。", "管理员可查看最近登录记录，辅助账号管理。"]),
          ("题目与测试用例", ["管理员可新建、编辑和删除题目。", "测试用例支持批量编辑、追加、删除和提交，并在独立确认操作后删除题目。 "])], None, None),
        ("安全、退出与维护", "软件通过认证、权限控制、隔离执行和配置外置降低常见安全风险。",
         [("安全措施", ["Spring Security与JWT保护需要认证的接口。", "Docker隔离用户代码编译运行；RabbitMQ实现任务解耦。", "图片存储错误对外返回通用提示，避免泄露存储端点和存储桶信息。", "真实配置通过环境变量管理，不提交到Git仓库。"]),
          ("退出与维护", ["用户选择退出后清除本地身份状态并返回未登录界面。", "升级版本前备份数据库和对象存储，并验证前后端测试与构建。 "])], None, None),
    ]
    for idx, (title, intro, sections, image, caption) in enumerate(pages, 1):
        add_manual_page(doc, idx, title, intro, sections, image, caption)
    path = OUT / "02-Blog and OJ V1.0软件操作说明书.docx"
    doc.save(path)
    return path


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--documents", choices=("basic", "manual", "all"), default="all")
    args = parser.parse_args()
    OUT.mkdir(parents=True, exist_ok=True)
    generated = []
    if args.documents in {"basic", "all"}:
        generated.extend((build_info_form(), build_guide(), build_pending()))
    if args.documents in {"manual", "all"}:
        generated.append(build_manual())
    for path in generated:
        print(path.relative_to(ROOT))


if __name__ == "__main__":
    main()
