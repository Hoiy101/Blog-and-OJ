#!/usr/bin/env python3
from __future__ import annotations

import hashlib
import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "copyright-application"
INVENTORY = json.loads((OUT / "qa" / "source-inventory.json").read_text(encoding="utf-8"))
DOCX_PATH = OUT / "03-Blog and OJ V1.0源程序鉴别材料.docx"
SELECTION_PATH = OUT / "qa" / "source-selection.json"
LINES_PER_PAGE = 50
TOTAL_PAGES = 60
SEGMENT_LINES = 1500


def font(run, name="Menlo", size=6.8, bold=False, color=None) -> None:
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), name)
    rpr.rFonts.set(qn("w:hAnsi"), name)
    rpr.rFonts.set(qn("w:eastAsia"), "STHeiti")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_page_number_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend((fld_char1, instr_text, fld_char2))
    font(run, "STHeiti", 8, color="666666")


def load_canonical_lines() -> list[dict]:
    safe = [item for item in INVENTORY["files"] if not item["sensitive_matches"]]
    safe.sort(key=lambda item: (
        0 if item["path"].startswith("web/src/") else
        1 if "/backend/src/main/java/" in item["path"] else 2,
        item["path"],
    ))
    records = []
    for item in safe:
        path = ROOT / item["path"]
        lines = path.read_text(encoding="utf-8", errors="replace").replace("\r\n", "\n").splitlines()
        for number, text in enumerate(lines, 1):
            records.append({"path": item["path"], "line": number, "text": text.expandtabs(4)})
    if len(records) < SEGMENT_LINES * 2:
        raise RuntimeError(f"Safe source has only {len(records)} lines; 3000 required")
    return records


def choose_lines(records: list[dict]) -> list[dict]:
    selected = records[:SEGMENT_LINES] + records[-SEGMENT_LINES:]
    if len(selected) != SEGMENT_LINES * 2:
        raise AssertionError("Selection must contain exactly 3000 source lines")
    for index, record in enumerate(selected, 1):
        record = record.copy()
        record["deposit_line"] = index
        record["page"] = (index - 1) // LINES_PER_PAGE + 1
        record["page_line"] = (index - 1) % LINES_PER_PAGE + 1
        selected[index - 1] = record
    return selected


def build_doc(selected: list[dict]) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.48)
    section.right_margin = Inches(0.48)
    section.header_distance = Inches(0.22)
    section.footer_distance = Inches(0.22)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(header.add_run("Blog and OJ V1.0 源程序鉴别材料"), "STHeiti", 8.5, bold=True, color="333333")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(footer.add_run("第 "), "STHeiti", 8, color="666666")
    add_page_number_field(footer)
    font(footer.add_run(" 页 / 共 60 页"), "STHeiti", 8, color="666666")
    normal = doc.styles["Normal"]
    normal.font.name = "Menlo"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Menlo")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Menlo")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "STHeiti")
    normal.font.size = Pt(6.8)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = Pt(8.2)
    for page in range(1, TOTAL_PAGES + 1):
        if page > 1:
            doc.add_page_break()
        chunk = selected[(page - 1) * LINES_PER_PAGE:page * LINES_PER_PAGE]
        paths = []
        for row in chunk:
            if not paths or paths[-1] != row["path"]:
                paths.append(row["path"])
        meta = doc.add_paragraph()
        meta.paragraph_format.space_after = Pt(3)
        meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
        label = f"第{page:02d}页 | 行{chunk[0]['deposit_line']:04d}-{chunk[-1]['deposit_line']:04d} | "
        font(meta.add_run(label), "STHeiti", 7.6, bold=True, color="2E74B5")
        path_text = paths[0] if len(paths) == 1 else f"{paths[0]} ... {paths[-1]}"
        font(meta.add_run(path_text), "STHeiti", 6.8, color="555555")
        code = doc.add_paragraph()
        code.paragraph_format.space_before = Pt(0)
        code.paragraph_format.space_after = Pt(0)
        code.paragraph_format.line_spacing = Pt(8.2)
        for idx, row in enumerate(chunk):
            display = row["text"].replace("\u0000", "")
            if len(display) > 132:
                display = display[:129] + "..."
            run = code.add_run(f"{row['deposit_line']:04d}  {display}")
            font(run)
            if idx < len(chunk) - 1:
                run.add_break()
    doc.core_properties.title = "Blog and OJ V1.0源程序鉴别材料"
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.save(DOCX_PATH)


def main() -> None:
    records = load_canonical_lines()
    selected = choose_lines(records)
    selection = {
        "software": "Blog and OJ V1.0",
        "deposit_method": "普通交存：规范化源码序列前1500行和后1500行",
        "pages": TOTAL_PAGES,
        "lines_per_page": LINES_PER_PAGE,
        "total_lines": len(selected),
        "selection_sha256": hashlib.sha256("\n".join(
            f"{row['path']}:{row['line']}:{row['text']}" for row in selected
        ).encode("utf-8")).hexdigest(),
        "lines": selected,
    }
    SELECTION_PATH.parent.mkdir(parents=True, exist_ok=True)
    SELECTION_PATH.write_text(json.dumps(selection, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    build_doc(selected)
    print(f"Selected {len(selected)} lines from {len(records)} safe canonical lines")
    print(DOCX_PATH.relative_to(ROOT))
    print(SELECTION_PATH.relative_to(ROOT))


if __name__ == "__main__":
    main()
